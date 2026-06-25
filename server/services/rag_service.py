"""
RAG 检索服务。

负责：
1. 初始化并持久化 Chroma 向量库；
2. 解析多种格式文档（PDF/TXT/Markdown/DOCX）并切片；
3. 将切片写入向量库 / 按文档删除向量；
4. 相似度检索，供 Agent 工具调用。
"""

import logging
from pathlib import Path

from langchain_chroma import Chroma
from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from server.core.config import CHROMA_DIR, settings
from server.services.llm import get_embeddings

logger = logging.getLogger(__name__)

# 走 MinerU 高保真解析的文件类型（其余如 txt/md 本身即纯文本，直接读取）
_MINERU_TYPES = {"pdf", "docx"}

# 向量库集合名称
COLLECTION_NAME = "enterprise_kb"

# ===== 检索链路参数 =====
# 向量召回数量
_VECTOR_TOPN = 10
# BM25 召回数量
_BM25_TOPN = 10
# RRF 融合常数（越大，名次差异影响越小）
_RRF_C = 60
# 送入重排的候选数量
_RERANK_CANDIDATES = 10

# 单例缓存
_vectorstore: Chroma | None = None
# BM25 检索器缓存（语料变更时置空以便重建）
_bm25_retriever: BM25Retriever | None = None


def get_vectorstore() -> Chroma:
    """获取（懒加载）Chroma 向量库单例。"""
    global _vectorstore
    if _vectorstore is None:
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        _vectorstore = Chroma(
            collection_name=COLLECTION_NAME,
            embedding_function=get_embeddings(),
            persist_directory=str(CHROMA_DIR),
        )
    return _vectorstore


def _read_pdf(path: Path) -> str:
    """读取 PDF 文件全部文本。"""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    """读取 DOCX 文件全部文本。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_text(path: Path) -> str:
    """读取纯文本/Markdown 文件。"""
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_with_mineru(path: Path) -> str:
    """
    使用 MinerU 的 LangChain Loader 高保真解析文档（precision 模式）。

    相比本地 pypdf/python-docx，MinerU 能更好地处理复杂排版、表格、公式与扫描件（OCR），
    并输出 Markdown 文本，更利于后续切片与检索。需在 .env 配置 MINERU_TOKEN。
    返回拼接后的纯文本（Markdown）。
    """
    from langchain_mineru import MinerULoader

    loader = MinerULoader(
        source=str(path),
        mode="precision",
        token=settings.mineru_token or None,
        split_pages=False,
        ocr=True,
        formula=True,
        table=True,
    )
    docs = loader.load()
    return "\n\n".join(d.page_content for d in docs if d.page_content)


def _read_local(path: Path, ftype: str) -> str:
    """本地解析兜底：MinerU 不可用时退回 pypdf/python-docx/纯文本。"""
    if ftype == "pdf":
        return _read_pdf(path)
    if ftype == "docx":
        return _read_docx(path)
    return _read_text(path)


def parse_file(file_path: str, file_type: str) -> str:
    """
    根据文件类型解析出纯文本内容。

    - pdf/docx：优先用 MinerU 高保真解析（失败则本地兜底）；
    - txt/md 及其它：直接按纯文本读取。
    """
    path = Path(file_path)
    ftype = file_type.lower()

    if ftype in _MINERU_TYPES and settings.mineru_token:
        try:
            text = _read_with_mineru(path)
            if text.strip():
                return text
            logger.warning("MinerU 解析结果为空，回退本地解析：%s", path.name)
        except Exception as exc:  # noqa: BLE001 - 解析失败回退本地方案，保证上传不中断
            logger.warning("MinerU 解析失败（%s），回退本地解析：%s", exc, path.name)
        return _read_local(path, ftype)

    return _read_local(path, ftype)


def ingest_document(doc_id: int, title: str, file_path: str, file_type: str) -> int:
    """
    解析文档并切片写入向量库。

    每个切片都带上 doc_id 与 title 元数据，便于检索溯源与删除。
    返回写入的切片数量。
    """
    text = parse_file(file_path, file_type)
    if not text.strip():
        return 0

    # 递归字符切分器：中文场景下块大小 500、重叠 100
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
    )
    chunks = splitter.split_text(text)

    documents: list[LCDocument] = []
    ids: list[str] = []
    for index, chunk in enumerate(chunks):
        documents.append(
            LCDocument(
                page_content=chunk,
                metadata={"doc_id": str(doc_id), "title": title, "chunk_index": index},
            )
        )
        ids.append(f"doc_{doc_id}_chunk_{index}")

    if documents:
        vectorstore = get_vectorstore()
        vectorstore.add_documents(documents=documents, ids=ids)
        _invalidate_bm25()
    return len(documents)


def delete_document_vectors(doc_id: int) -> None:
    """删除指定文档对应的全部向量切片。"""
    vectorstore = get_vectorstore()
    vectorstore.delete(where={"doc_id": str(doc_id)})
    _invalidate_bm25()


# ===========================================================================
# 检索链路：向量召回 + BM25 召回 -> RRF 融合 -> Qwen 重排
# ===========================================================================


def _invalidate_bm25() -> None:
    """语料发生变更后置空 BM25 缓存，下次检索时重建。"""
    global _bm25_retriever
    _bm25_retriever = None


def _chunk_key(doc: LCDocument) -> tuple[str, str]:
    """切片唯一键：(doc_id, chunk_index)，用于跨召回路去重与融合。"""
    meta = doc.metadata or {}
    return (str(meta.get("doc_id", "")), str(meta.get("chunk_index", "")))


def _get_bm25_retriever() -> BM25Retriever | None:
    """
    懒加载 BM25 检索器：从 Chroma 取出全部切片，用 jieba 分词构建。

    中文需先分词，否则 BM25 会把整句当作一个 token。语料为空或构建失败返回 None。
    """
    global _bm25_retriever
    if _bm25_retriever is not None:
        return _bm25_retriever
    try:
        import jieba

        data = get_vectorstore().get(include=["documents", "metadatas"])
        texts = data.get("documents") or []
        metadatas = data.get("metadatas") or []
        if not texts:
            return None
        docs = [
            LCDocument(page_content=t, metadata=m or {})
            for t, m in zip(texts, metadatas)
        ]
        retriever = BM25Retriever.from_documents(
            docs, preprocess_func=lambda text: jieba.lcut(text)
        )
        retriever.k = _BM25_TOPN
        _bm25_retriever = retriever
        return _bm25_retriever
    except Exception as exc:  # noqa: BLE001 - BM25 不可用时降级为仅向量检索
        logger.warning("BM25 构建失败，跳过关键词召回：%s", exc)
        return None


def _rrf_fuse(
    ranked_lists: list[list[LCDocument]], c: int = _RRF_C
) -> list[LCDocument]:
    """
    Reciprocal Rank Fusion：按各路召回的名次融合，score = Σ 1/(c + rank)。

    rank 从 0 开始计；同一切片在多路命中则分数累加，最终按总分降序。
    """
    scores: dict[tuple[str, str], float] = {}
    doc_map: dict[tuple[str, str], LCDocument] = {}
    for ranked in ranked_lists:
        for rank, doc in enumerate(ranked):
            key = _chunk_key(doc)
            scores[key] = scores.get(key, 0.0) + 1.0 / (c + rank)
            doc_map.setdefault(key, doc)
    ordered_keys = sorted(scores, key=lambda kk: scores[kk], reverse=True)
    return [doc_map[k] for k in ordered_keys]


def _rerank(query: str, docs: list[LCDocument], top_k: int) -> list[LCDocument]:
    """用 DashScope(Qwen) rerank 模型对候选片段精排，取 Top-k；失败则返回原顺序前 k 个。"""
    if not docs:
        return []
    try:
        import dashscope

        resp = dashscope.TextReRank.call(
            model=settings.rerank_model,
            query=query,
            documents=[d.page_content for d in docs],
            top_n=top_k,
            return_documents=False,
            api_key=settings.dashscope_api_key,
        )
        results = (resp.output or {}).get("results") if hasattr(resp, "output") else None
        if not results:
            logger.warning("rerank 无返回，回退融合顺序")
            return docs[:top_k]
        reranked = [docs[item["index"]] for item in results if item.get("index") is not None]
        return reranked[:top_k] or docs[:top_k]
    except Exception as exc:  # noqa: BLE001 - rerank 失败时回退融合顺序
        logger.warning("rerank 失败，回退融合顺序：%s", exc)
        return docs[:top_k]


def search(query: str, k: int = 4) -> list[LCDocument]:
    """
    混合检索链路：向量召回 + BM25 关键词召回 -> RRF 融合 -> Qwen 重排，返回 Top-k 切片。

    每个环节均带降级：BM25 不可用则仅用向量；rerank 不可用则用融合顺序；
    始终保证在向量库可用时能返回结果。
    """
    vectorstore = get_vectorstore()

    # 1) 向量召回
    try:
        vector_docs = vectorstore.similarity_search(query, k=_VECTOR_TOPN)
    except Exception as exc:  # noqa: BLE001
        logger.warning("向量检索失败：%s", exc)
        vector_docs = []

    # 2) BM25 关键词召回（可降级）
    bm25_docs: list[LCDocument] = []
    bm25 = _get_bm25_retriever()
    if bm25 is not None:
        try:
            bm25_docs = bm25.invoke(query)
        except Exception as exc:  # noqa: BLE001
            logger.warning("BM25 检索失败：%s", exc)
            bm25_docs = []

    # 3) RRF 融合（仅一路有结果时即等价于该路顺序）
    ranked_lists = [lst for lst in (vector_docs, bm25_docs) if lst]
    if not ranked_lists:
        return []
    fused = _rrf_fuse(ranked_lists)

    # 4) 取候选送入重排，返回 Top-k
    candidates = fused[:_RERANK_CANDIDATES]
    return _rerank(query, candidates, top_k=k)
