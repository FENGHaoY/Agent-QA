"""生成"易混淆"评测语料（用于检验混合检索 + Rerank 相对纯向量检索的增益）。

设计思路：构造 4 个"同族"文档族，族内文档主题/结构高度相似，
仅靠 专有名词 / 编号 / 版本号 / 金额 等关键 token 区分。
纯向量检索对这类"语义近似、实体不同"的文档容易混淆，
而 BM25 关键词召回 + Rerank 重排能凭精确 token 命中正确文档。

注意：
- 输出到独立目录 eval_docs/，不与已有的 test_docs/ 混合。
- 入库时后端取「文件名去扩展名」作为 title，因此把区分码写进了文件名。

运行：uv run python scripts/gen_eval_docs.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "eval_docs"
CN_FONT = "STSong-Light"

Doc = tuple[str, str, list[str]]


# ===========================================================================
# 族 A：数据中心运维手册（docx ×3）—— 结构一致，仅机房编号/设备型号/IP/联系人不同
# ===========================================================================
DOCX_DOCS: list[Doc] = [
    (
        "数据中心运维手册-IDC-BJ-01.docx",
        "数据中心运维手册-IDC-BJ-01",
        [
            "一、机房概况",
            "本机房编号 IDC-BJ-01，位于北京亦庄经济技术开发区科创五街 18 号，"
            "总机柜数 320 个，设计 PUE 1.35，等级为 T3+。机房主管：王海。",
            "二、供配电系统",
            "采用双路市电输入，配备康明斯 C1100 柴油发电机作为后备电源，"
            "UPS 采用 APC Symmetra PX 80kW，后备时间不少于 30 分钟。",
            "三、制冷系统",
            "采用海洛斯（Hiross）ASD 系列精密空调，N+1 冗余，机房温度维持在 22±2℃，"
            "相对湿度 45%~55%。",
            "四、网络架构",
            "核心交换机为华为 CloudEngine CE12808，IP 核心网段为 10.10.0.0/16，"
            "出口带宽 20Gbps，BGP 多线接入。",
            "五、门禁与安防",
            "门禁系统采用 ZKTeco ProID20 指纹+刷卡双重认证，全区域 7×24 视频监控，"
            "录像保存不少于 90 天。",
            "六、值班与应急",
            "运维值班电话 010-8888-0001，发生重大故障应在 15 分钟内上报主管王海，"
            "并启动 IDC-BJ-01 应急预案。",
        ],
    ),
    (
        "数据中心运维手册-IDC-SH-02.docx",
        "数据中心运维手册-IDC-SH-02",
        [
            "一、机房概况",
            "本机房编号 IDC-SH-02，位于上海金桥出口加工区张江路 99 号，"
            "总机柜数 480 个，设计 PUE 1.30，等级为 T3+。机房主管：李雷。",
            "二、供配电系统",
            "采用双路市电输入，配备卡特彼勒 C1250 柴油发电机作为后备电源，"
            "UPS 采用华为 UPS5000-E 120kW，后备时间不少于 30 分钟。",
            "三、制冷系统",
            "采用维谛（Vertiv）Liebert PEX 精密空调，N+1 冗余，机房温度维持在 22±2℃，"
            "相对湿度 45%~55%。",
            "四、网络架构",
            "核心交换机为锐捷 RG-N18010，IP 核心网段为 10.20.0.0/16，"
            "出口带宽 40Gbps，BGP 多线接入。",
            "五、门禁与安防",
            "门禁系统采用海康威视 DS-K 人脸+刷卡双重认证，全区域 7×24 视频监控，"
            "录像保存不少于 90 天。",
            "六、值班与应急",
            "运维值班电话 021-8888-0002，发生重大故障应在 15 分钟内上报主管李雷，"
            "并启动 IDC-SH-02 应急预案。",
        ],
    ),
    (
        "数据中心运维手册-IDC-GZ-03.docx",
        "数据中心运维手册-IDC-GZ-03",
        [
            "一、机房概况",
            "本机房编号 IDC-GZ-03，位于广州天河区高普路 6 号，"
            "总机柜数 260 个，设计 PUE 1.40，等级为 T3。机房主管：陈明。",
            "二、供配电系统",
            "采用双路市电输入，配备沃尔沃 TAD1640 柴油发电机作为后备电源，"
            "UPS 采用科华 KR33 60kW，后备时间不少于 30 分钟。",
            "三、制冷系统",
            "采用佳力图（CyberMate）精密空调，N+1 冗余，机房温度维持在 22±2℃，"
            "相对湿度 45%~55%。",
            "四、网络架构",
            "核心交换机为思科 Nexus 9508，IP 核心网段为 10.30.0.0/16，"
            "出口带宽 10Gbps，BGP 多线接入。",
            "五、门禁与安防",
            "门禁系统采用大华 ASI 系列刷卡认证，全区域 7×24 视频监控，"
            "录像保存不少于 90 天。",
            "六、值班与应急",
            "运维值班电话 020-8888-0003，发生重大故障应在 15 分钟内上报主管陈明，"
            "并启动 IDC-GZ-03 应急预案。",
        ],
    ),
]


# ===========================================================================
# 族 C：灵犀智能客服平台版本发布说明（txt ×3）—— 同一产品不同版本
# ===========================================================================
TXT_DOCS: list[Doc] = [
    (
        "灵犀智能客服平台-v3.1发布说明.txt",
        "灵犀智能客服平台-v3.1发布说明",
        [
            "一、版本信息",
            "产品名称：灵犀（LingXi）智能客服平台。版本号：v3.1.0。"
            "发布日期：2024-03-10。对应需求编号：LX-REQ-310。升级包：lingxi-3.1.0.tar.gz。",
            "二、新增功能",
            "1. 新增多轮对话上下文记忆能力，支持引用前文进行追问。",
            "2. 新增坐席工作台快捷回复模板管理。",
            "三、优化项",
            "优化了知识库检索响应速度，平均检索耗时由 800ms 降至 450ms。",
            "四、修复缺陷",
            "修复 BUG-3105：高并发场景下知识库检索偶发超时。"
            "修复 BUG-3110：导出会话记录中文乱码问题。",
            "五、已知问题",
            "在 IE11 浏览器下富文本编辑器样式异常，建议使用 Chrome。",
            "六、升级指引",
            "从 v3.0 升级请先备份配置文件 config.yaml，再执行 ./upgrade.sh lingxi-3.1.0.tar.gz。",
        ],
    ),
    (
        "灵犀智能客服平台-v3.2发布说明.txt",
        "灵犀智能客服平台-v3.2发布说明",
        [
            "一、版本信息",
            "产品名称：灵犀（LingXi）智能客服平台。版本号：v3.2.0。"
            "发布日期：2024-06-12。对应需求编号：LX-REQ-320。升级包：lingxi-3.2.0.tar.gz。",
            "二、新增功能",
            "1. 新增工单自动分类能力，可按业务类型自动路由至对应坐席组。",
            "2. 新增满意度评价短信回访。",
            "三、优化项",
            "优化了坐席并发会话的内存占用，单实例最大并发会话数由 200 提升至 500。",
            "四、修复缺陷",
            "修复 BUG-3208：高并发下部分会话状态丢失。"
            "修复 BUG-3215：工单附件上传大于 10MB 时失败。",
            "五、已知问题",
            "工单自动分类对极短文本（少于 5 字）准确率偏低。",
            "六、升级指引",
            "从 v3.1 升级请先备份配置文件 config.yaml，再执行 ./upgrade.sh lingxi-3.2.0.tar.gz。",
        ],
    ),
    (
        "灵犀智能客服平台-v4.0发布说明.txt",
        "灵犀智能客服平台-v4.0发布说明",
        [
            "一、版本信息",
            "产品名称：灵犀（LingXi）智能客服平台。版本号：v4.0.0。"
            "发布日期：2024-09-20。对应需求编号：LX-REQ-400。升级包：lingxi-4.0.0.tar.gz。",
            "二、新增功能",
            "1. 新增基于大模型的 RAG 智能问答，支持企业知识库自动检索作答。",
            "2. 整体架构由单体升级为微服务，支持横向扩缩容。",
            "三、优化项",
            "重构了消息推送链路，端到端时延降低约 40%。",
            "四、修复缺陷",
            "修复 BUG-4012：长时间运行后网关进程内存泄漏。",
            "五、已知问题",
            "v4.0 不兼容 v3.x 的旧版配置，需使用迁移工具 migrate-4.0 转换。",
            "六、升级指引",
            "v4.0 为重大版本，需停机升级；请先使用 migrate-4.0 转换配置，"
            "再执行 ./upgrade.sh lingxi-4.0.0.tar.gz。",
        ],
    ),
]


# ===========================================================================
# 族 D：采购订单（pdf ×3）—— 结构一致，仅订单号/供应商/设备/金额不同
# ===========================================================================
PDF_DOCS: list[Doc] = [
    (
        "采购订单-PO-2024-0312.pdf",
        "采购订单-PO-2024-0312",
        [
            "一、订单信息",
            "采购订单编号：PO-2024-0312。下单日期：2024-03-28。需求部门：基础架构部。"
            "经办人：刘洋。",
            "二、供应商信息",
            "供应商名称：鼎信科技有限公司。纳税人识别号：91110302MA01A12X3K。"
            "联系人：钱进，电话 138-0000-0312。",
            "三、采购明细",
            "采购设备：Dell PowerEdge R750 机架式服务器，数量 10 台，单价 ¥12,800。",
            "四、金额与付款",
            "订单总金额：¥128,000（含税）。付款方式：预付 30%，验收合格后付清尾款。",
            "五、交付与验收",
            "约定交付日期：2024-04-15。收货地点：IDC-BJ-01，收货人：王海。"
            "到货后 5 个工作日内完成验收。",
            "六、售后服务",
            "整机质保 3 年，提供 7×24 原厂支持，故障 4 小时响应。",
        ],
    ),
    (
        "采购订单-PO-2024-0418.pdf",
        "采购订单-PO-2024-0418",
        [
            "一、订单信息",
            "采购订单编号：PO-2024-0418。下单日期：2024-04-30。需求部门：网络部。"
            "经办人：刘洋。",
            "二、供应商信息",
            "供应商名称：华彩网络科技有限公司。纳税人识别号：91310115MA1G88K7P2。"
            "联系人：孙立，电话 139-0000-0418。",
            "三、采购明细",
            "采购设备：华为 CloudEngine S5732 千兆交换机，数量 8 台，单价 ¥7,000。",
            "四、金额与付款",
            "订单总金额：¥56,000（含税）。付款方式：货到验收合格后一次性付清。",
            "五、交付与验收",
            "约定交付日期：2024-05-20。收货地点：IDC-SH-02，收货人：李雷。"
            "到货后 5 个工作日内完成验收。",
            "六、售后服务",
            "整机质保 1 年，提供工作日 5×8 远程支持。",
        ],
    ),
    (
        "采购订单-PO-2024-0529.pdf",
        "采购订单-PO-2024-0529",
        [
            "一、订单信息",
            "采购订单编号：PO-2024-0529。下单日期：2024-05-29。需求部门：存储与备份组。"
            "经办人：刘洋。",
            "二、供应商信息",
            "供应商名称：鼎信科技有限公司。纳税人识别号：91110302MA01A12X3K。"
            "联系人：钱进，电话 138-0000-0529。",
            "三、采购明细",
            "采购设备：NetApp AFF A400 全闪存存储阵列，数量 2 台，单价 ¥160,000。",
            "四、金额与付款",
            "订单总金额：¥320,000（含税）。付款方式：分三期支付，到货付 50%，"
            "验收付 40%，质保期满付 10%。",
            "五、交付与验收",
            "约定交付日期：2024-06-30。收货地点：IDC-GZ-03，收货人：陈明。"
            "到货后 10 个工作日内完成验收。",
            "六、售后服务",
            "整机质保 5 年，提供 7×24 原厂支持，故障 2 小时响应。",
        ],
    ),
]


# ===========================================================================
# 族 B：微服务接口文档（md ×3）—— 结构一致，仅服务名/端口/路径/错误码不同
# ===========================================================================
MD_DOCS: list[tuple[str, str]] = [
    (
        "微服务接口文档-订单服务ORD-SVC.md",
        """# 微服务接口文档-订单服务 ORD-SVC

## 服务简介

订单服务（服务标识 ORD-SVC）负责订单的创建、查询、取消与状态流转，
部署端口 8081，基础路径 `/api/v1/order`，代码仓库 `git@git.example.com:svc/order.git`。

## 鉴权方式

所有接口需在请求头携带 `Authorization: Bearer <token>`，令牌由统一认证中心签发。

## 主要接口

- 创建订单：`POST /api/v1/order/orders`
- 查询订单：`GET /api/v1/order/orders/{orderId}`
- 取消订单：`POST /api/v1/order/orders/{orderId}/cancel`

## 错误码

- `ORD-40001`：请求参数校验失败
- `ORD-40901`：库存不足，无法下单
- `ORD-50000`：订单服务内部错误

## 限流与超时

默认限流 200 QPS，单请求超时时间 3000ms，超时返回 `ORD-50400`。

## 负责人

服务负责人：赵敏，邮箱 zhaomin@example.com。
""",
    ),
    (
        "微服务接口文档-支付服务PAY-SVC.md",
        """# 微服务接口文档-支付服务 PAY-SVC

## 服务简介

支付服务（服务标识 PAY-SVC）负责发起支付、退款与对账，
部署端口 8082，基础路径 `/api/v1/pay`，代码仓库 `git@git.example.com:svc/pay.git`。

## 鉴权方式

所有接口需在请求头携带 `Authorization: Bearer <token>`，并额外校验签名 `X-Sign`。

## 主要接口

- 发起支付：`POST /api/v1/pay/payments`
- 申请退款：`POST /api/v1/pay/refunds`
- 支付查询：`GET /api/v1/pay/payments/{payId}`

## 错误码

- `PAY-40010`：支付金额非法
- `PAY-50001`：第三方支付渠道异常
- `PAY-40310`：签名校验失败

## 限流与超时

默认限流 100 QPS，单请求超时时间 5000ms，超时返回 `PAY-50400`。

## 负责人

服务负责人：孙强，邮箱 sunqiang@example.com。
""",
    ),
    (
        "微服务接口文档-库存服务INV-SVC.md",
        """# 微服务接口文档-库存服务 INV-SVC

## 服务简介

库存服务（服务标识 INV-SVC）负责库存的查询、扣减与回补，
部署端口 8083，基础路径 `/api/v1/inventory`，代码仓库 `git@git.example.com:svc/inventory.git`。

## 鉴权方式

所有接口需在请求头携带 `Authorization: Bearer <token>`，内部调用走 mTLS。

## 主要接口

- 扣减库存：`POST /api/v1/inventory/deductions`
- 回补库存：`POST /api/v1/inventory/restorations`
- 查询库存：`GET /api/v1/inventory/stocks/{sku}`

## 错误码

- `INV-40020`：SKU 不存在
- `INV-60001`：库存不足，扣减失败
- `INV-50000`：库存服务内部错误

## 限流与超时

默认限流 300 QPS，单请求超时时间 2000ms，超时返回 `INV-50400`。

## 负责人

服务负责人：周婷，邮箱 zhouting@example.com。
""",
    ),
]


def write_txt(path: Path, title: str, paragraphs: list[str]) -> None:
    lines = [title, ""]
    lines.extend(paragraphs)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(str(path))


def write_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont(CN_FONT))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CNTitle", parent=styles["Title"], fontName=CN_FONT, fontSize=18, leading=24
    )
    body_style = ParagraphStyle(
        "CNBody", parent=styles["Normal"], fontName=CN_FONT, fontSize=11, leading=20
    )
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    story: list = [Paragraph(title, title_style), Spacer(1, 8 * mm)]
    for para in paragraphs:
        story.append(Paragraph(para, body_style))
        story.append(Spacer(1, 3 * mm))
    doc.build(story)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    count = 0

    for filename, title, paragraphs in DOCX_DOCS:
        write_docx(OUTPUT_DIR / filename, title, paragraphs)
        count += 1
        print(f"已生成 {filename}")

    for filename, title, paragraphs in TXT_DOCS:
        write_txt(OUTPUT_DIR / filename, title, paragraphs)
        count += 1
        print(f"已生成 {filename}")

    for filename, title, paragraphs in PDF_DOCS:
        write_pdf(OUTPUT_DIR / filename, title, paragraphs)
        count += 1
        print(f"已生成 {filename}")

    for filename, content in MD_DOCS:
        (OUTPUT_DIR / filename).write_text(content, encoding="utf-8")
        count += 1
        print(f"已生成 {filename}")

    print(f"\n完成：共生成 {count} 个易混淆评测文档，输出目录：{OUTPUT_DIR}")


if __name__ == "__main__":
    main()
